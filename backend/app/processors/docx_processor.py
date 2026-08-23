"""
DOCXProcessor — FileProcessor implementation for DOCX files.

For POC, classifies all content as UNSTRUCTURED → RAG.
Architecture supports future extension to detect embedded tables
and route them to DATASET_QUERY.
"""

from __future__ import annotations

import logging
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.errors.ingestion_errors import FileProcessingError
from app.processors.protocol import (
    DetectedRegion,
    InspectionResult,
    NormalizedDataset,
    ValidationWarning,
)

logger = logging.getLogger(__name__)

# Maximum paragraphs to include in content_sample
_MAX_SAMPLE_PARAGRAPHS = 20

# Mapping of DOCX heading styles to Markdown heading levels
_HEADING_STYLE_MAP: dict[str, str] = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Heading 5": "##### ",
    "Heading 6": "###### ",
}


class DOCXProcessor:
    """FileProcessor for DOCX files.

    For POC, classifies all content as UNSTRUCTURED → RAG.
    Architecture supports future extension to detect embedded tables.
    """

    def can_process(self, file_type: str) -> bool:
        """Return True if this processor handles the given file type.

        Args:
            file_type: File extension or type identifier.

        Returns:
            True for "docx".
        """
        return file_type.lower() == "docx"

    async def inspect(self, file_path: str) -> InspectionResult:
        """Inspect DOCX and produce a single UNSTRUCTURED region.

        Extracts paragraphs and table cell text using python-docx.
        Converts Heading 1-6 styles to Markdown format. Excludes headers,
        footers, and embedded objects.

        Args:
            file_path: Path to the DOCX file on disk.

        Returns:
            InspectionResult with a single region representing the full document.

        Raises:
            FileProcessingError: If the file is corrupted, unreadable, or
                password-protected.
        """
        path = Path(file_path)
        file_name = path.name

        if not path.exists():
            raise FileProcessingError(
                file_name=file_name,
                message=f"DOCX file not found: {file_path}",
            )

        try:
            doc = DocxDocument(file_path)
        except PackageNotFoundError as exc:
            raise FileProcessingError(
                file_name=file_name,
                message=f"Failed to open DOCX file (corrupted or invalid): {file_name}",
                detail=str(exc),
            ) from exc
        except Exception as exc:
            # Catch password-protected or other unreadable formats
            error_msg = str(exc).lower()
            if "password" in error_msg or "encrypted" in error_msg:
                raise FileProcessingError(
                    file_name=file_name,
                    message=f"DOCX file is password-protected: {file_name}",
                    detail=str(exc),
                ) from exc
            raise FileProcessingError(
                file_name=file_name,
                message=f"Failed to open DOCX file: {file_name}",
                detail=str(exc),
            ) from exc

        try:
            text_lines = self._extract_text_lines(doc)
        except Exception as exc:
            raise FileProcessingError(
                file_name=file_name,
                message=f"Error processing DOCX file: {file_name}",
                detail=str(exc),
            ) from exc

        # Build raw text
        raw_text = "\n".join(text_lines)

        # Build metadata
        metadata: dict[str, str] = {}
        core_props = doc.core_properties
        if core_props.title:
            metadata["title"] = core_props.title
        if core_props.author:
            metadata["author"] = core_props.author
        if core_props.created:
            metadata["creation_date"] = str(core_props.created)

        # Build content_sample from first paragraphs (non-empty lines)
        non_empty_lines = [line for line in text_lines if line.strip()]
        content_sample: list[list[str]] = [
            [line] for line in non_empty_lines[:_MAX_SAMPLE_PARAGRAPHS]
        ]

        # If no text content, return an empty region
        region = DetectedRegion(
            region_id=str(uuid4()),
            sheet_name="document",
            start_row=1,
            end_row=max(len(text_lines), 1),
            start_column=0,
            end_column=0,
            header_row=None,
            content_sample=content_sample,
            row_count=len(text_lines),
            column_count=1,
            raw_text=raw_text,
        )

        return InspectionResult(
            file_name=file_name,
            file_type="docx",
            regions=[region],
            metadata=metadata,
        )

    def _extract_text_lines(self, doc: DocxDocument) -> list[str]:
        """Extract text lines from paragraphs and tables.

        Converts heading styles to Markdown format. Includes table cell text
        but excludes headers, footers, and embedded objects.

        Args:
            doc: Parsed DOCX document.

        Returns:
            List of text lines with headings converted to Markdown.
        """
        lines: list[str] = []

        # Extract paragraphs with heading conversion
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                lines.append("")
                continue

            style_name = paragraph.style.name if paragraph.style else ""
            prefix = _HEADING_STYLE_MAP.get(style_name, "")
            lines.append(f"{prefix}{text}")

        # Extract table cell text
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    lines.append(" | ".join(row_texts))

        return lines

    async def extract(
        self, file_path: str, region: DetectedRegion | None = None
    ) -> NormalizedDataset:
        """Not applicable for unstructured DOCX content.

        DOCX content is routed to RAG pipeline, not structured extraction.

        Raises:
            NotImplementedError: Always — unstructured content is not applicable
                for structured extraction.
        """
        raise NotImplementedError(
            "DOCXProcessor does not support structured extraction. "
            "DOCX content is processed through the RAG pipeline."
        )

    def validate(self, normalized: NormalizedDataset) -> list[ValidationWarning]:
        """Validate a normalized dataset. Returns empty list for DOCX content.

        Args:
            normalized: The normalized dataset to validate.

        Returns:
            Empty list — DOCX content is unstructured and not validated this way.
        """
        return []
